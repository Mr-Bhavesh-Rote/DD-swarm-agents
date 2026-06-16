"""Word (.docx) export via python-docx (§9).

Headings per section, tables for structured data, and citations rendered as hyperlinked
[n] markers plus a References section. Matches the company report layout
(Executive Summary -> Ownership & Governance -> Operations Footprint -> Financials ->
Risk Issues -> Investment Considerations).
"""
from __future__ import annotations

import re
from typing import Any, Dict, List

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

from exporters.html import DISCLAIMER

_CITE_RE = re.compile(r"(\[\d+\])")
_CITE_ID_RE = re.compile(r"\[(\d+)\]")


def render_docx(report: Dict[str, Any], kind: str) -> bytes:
    import io

    doc = Document()
    src_by_id = {s["id"]: s for s in report.get("sources", [])}

    title = "RAW Research Output" if kind == "raw" else "Due-Diligence Report"
    doc.add_heading(f"{title} — {report.get('subject', '')}", level=0)

    meta = doc.add_paragraph()
    meta.add_run(f"Generated {report.get('generated_at', '')}").italic = True

    disc = doc.add_paragraph()
    disc.add_run(DISCLAIMER).italic = True

    if kind == "final":
        v = report.get("verification", {})
        p = doc.add_paragraph()
        p.add_run(
            f"Citation coverage: {v.get('citation_coverage', 0):.0%} · "
            f"Faithfulness: {v.get('faithfulness_score', 0):.0%}"
        ).bold = True
        _final_body(doc, report, src_by_id)
    else:
        _raw_body(doc, report)

    _references(doc, report.get("sources", []))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _final_body(doc, report: Dict[str, Any], src_by_id: Dict[int, Dict[str, Any]]) -> None:
    for sec in report.get("sections", []):
        doc.add_heading(sec.get("title", ""), level=1)
        for para in (sec.get("body_markdown", "") or "").split("\n\n"):
            if not para.strip():
                continue
            _add_paragraph_with_citations(doc, para.strip(), src_by_id)
        for t in sec.get("tables", []):
            _add_table(doc, t)


def _raw_body(doc, report: Dict[str, Any]) -> None:
    for ao in report.get("agent_outputs", []):
        doc.add_heading(f"{ao.get('role') or ao.get('agent', '')} ({ao.get('model', '')})", level=1)
        for para in (ao.get("narrative_markdown", "") or "").split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        findings = ao.get("findings", [])
        if findings:
            doc.add_heading("Findings", level=2)
            for f in findings:
                ids = f.get("source_ids") or f.get("source_urls") or []
                doc.add_paragraph(f"{f.get('claim', '')}  {ids}", style="List Bullet")


def _add_paragraph_with_citations(doc, text: str, src_by_id: Dict[int, Dict[str, Any]]) -> None:
    p = doc.add_paragraph()
    for token in _CITE_RE.split(text):
        m = _CITE_ID_RE.fullmatch(token)
        if m:
            cid = int(m.group(1))
            src = src_by_id.get(cid)
            if src:
                _add_hyperlink(p, src["url"], token)
                continue
        if token:
            p.add_run(token)


def _add_table(doc, t: Dict[str, Any]) -> None:
    cols = t.get("columns", [])
    if not cols:
        return
    if t.get("title"):
        doc.add_paragraph().add_run(t["title"]).bold = True
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Light Grid Accent 1"
    for i, c in enumerate(cols):
        table.rows[0].cells[i].text = str(c)
    for row in t.get("rows", []):
        cells = table.add_row().cells
        for i, val in enumerate(row[: len(cols)]):
            cells[i].text = str(val)


def _references(doc, sources: List[Dict[str, Any]]) -> None:
    doc.add_heading("References", level=1)
    for s in sources:
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"[{s['id']}] ")
        _add_hyperlink(p, s["url"], s.get("title") or s["url"])


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """Insert a real clickable hyperlink run into a python-docx paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1155CC")
    rpr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(u)
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
